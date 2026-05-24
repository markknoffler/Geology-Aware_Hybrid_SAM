#!/usr/bin/env python3
"""
Build conference_manuscript_tri_encoder_cfm.docx from:
  - conference_remotesensing_landslide.pdf (intro, related work, dataset blurb, references)
  - MODEL_ARCHITECTURE.md + ablation CSV + generated figures

Run from anywhere:
  python SAM/model_architecture_cfm_landseg/scripts/build_conference_docx.py
"""

from __future__ import annotations

import csv
import math
import re
import shutil
import sys
from pathlib import Path

try:
    from pypdf import PdfReader
except ImportError:
    import subprocess

    subprocess.check_call([sys.executable, "-m", "pip", "install", "pypdf", "-q"])
    from pypdf import PdfReader

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def repo_root() -> Path:
    """CSIR_NEIST repository root (parent of SAM/)."""
    return Path(__file__).resolve().parents[3]


PROPOSED_MODEL_ID = "tri_encoder_cfm_v2"


def extract_pdf_text(pdf_path: Path) -> str:
    r = PdfReader(str(pdf_path))
    parts = []
    for page in r.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)


def clean_header_lines(text: str) -> str:
    out_lines = []
    for ln in text.splitlines():
        s = ln.strip()
        if re.match(r"^\d+\s+S\.\s+Bhuyan", s):
            continue
        if "Title Suppressed Due to Excessive Length" in s:
            continue
        if s == "Keywords: Landslide modeling · Machine Learning · Segmentation · SAM · CNN.":
            continue
        out_lines.append(ln)
    return "\n".join(out_lines)


def slice_between(full: str, start_pat: str, end_pat: str | None, flags=re.DOTALL) -> str:
    m0 = re.search(start_pat, full, flags)
    if not m0:
        return ""
    start = m0.start()
    if end_pat:
        m1 = re.search(end_pat, full[m0.end() :], flags)
        if m1:
            return full[start : m0.end() + m1.start()].strip()
    return full[start:].strip()


def add_para(doc: Document, text: str, style: str | None = None) -> None:
    text = text.strip()
    if not text:
        return
    for block in re.split(r"\n\s*\n+", text):
        block = " ".join(block.split())
        if not block:
            continue
        p = doc.add_paragraph(block, style=style)


def _black() -> RGBColor:
    return RGBColor(0, 0, 0)


def _set_cell_edge_border(cell, edge: str, val: str, sz: str = "12", color: str = "000000") -> None:
    """edge: top|bottom|left|right — val 'single' or 'nil'."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tc_borders = tcPr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tcPr.append(tc_borders)
    for child in list(tc_borders):
        if child.tag.endswith(edge):
            tc_borders.remove(child)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), val)
    if val != "nil":
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")
    tc_borders.append(el)


def apply_journal_horizontal_rules_table(table) -> None:
    """Top rule, rule under header, bottom rule; no vertical rules."""
    rows = table.rows
    n = len(rows)
    if n == 0:
        return
    for ri, row in enumerate(rows):
        for cell in row.cells:
            _set_cell_edge_border(cell, "left", "nil")
            _set_cell_edge_border(cell, "right", "nil")
            if ri == 0:
                _set_cell_edge_border(cell, "top", "single", "18")
                _set_cell_edge_border(cell, "bottom", "single", "12")
            elif ri == n - 1:
                _set_cell_edge_border(cell, "top", "nil")
                _set_cell_edge_border(cell, "bottom", "single", "18")
            else:
                _set_cell_edge_border(cell, "top", "nil")
                _set_cell_edge_border(cell, "bottom", "nil")


def force_all_text_black(doc: Document) -> None:
    """Remove theme colours / hyperlinks blue — force RGB black on body runs."""
    black = _black()

    def paint_paragraph(p) -> None:
        for run in p.runs:
            run.font.color.rgb = black

    for p in doc.paragraphs:
        paint_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    paint_paragraph(p)
                for inner in cell.tables:
                    for irow in inner.rows:
                        for icell in irow.cells:
                            for ip in icell.paragraphs:
                                paint_paragraph(ip)
    for sec in doc.sections:
        for part in (sec.header, sec.footer):
            for p in part.paragraphs:
                paint_paragraph(p)


def add_figure(
    doc: Document,
    image_path: Path,
    caption: str,
    width_in: float = 6.2,
) -> None:
    if not image_path.is_file():
        p = doc.add_paragraph(f"[Missing image file: {image_path}]")
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = _black()
    else:
        doc.add_picture(str(image_path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = _black()
    doc.add_paragraph("")


def strip_first_numbered_heading(text: str) -> str:
    """Remove first line if it looks like '1 Introduction' / '2 Related Works' / '3.1 ...' from PDF extract."""
    lines = text.splitlines()
    if not lines:
        return text
    if re.match(r"^\d+(\.\d+)?\s+\S", lines[0].strip()):
        lines = lines[1:]
    return "\n".join(lines).strip()


def minimal_intro_edits(intro: str) -> str:
    intro = intro.replace(
        "the following work proposes a hybrid framework",
        "the following work originally proposed a hybrid framework",
    )
    if "Motivated by these issues" in intro:
        intro = re.sub(
            r"Motivated by these issues,.*detection\.",
            NEW_CONTRIBUTION_PARA.strip(),
            intro,
            count=1,
            flags=re.DOTALL,
        )
    return intro


def set_cell_shading(cell, fill_hex: str) -> None:
    """Light grey header shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _sort_rows_proposed_last(rows: list[dict[str, str]]) -> list[dict[str, str]]:
    pid = PROPOSED_MODEL_ID.lower()
    head = [r for r in rows if (r.get("model_id") or "").lower() != pid]
    tail = [r for r in rows if (r.get("model_id") or "").lower() == pid]
    return head + tail


def add_results_table(doc: Document, csv_path: Path, *, note_extra: str = "") -> None:
    rows_raw = list(csv.DictReader(csv_path.open(newline="", encoding="utf-8")))
    rows = _sort_rows_proposed_last(rows_raw)
    display_cols = [
        "model_id",
        "display_name",
        "best_epoch",
        "val_acc",
        "val_precision",
        "val_recall",
        "val_f1",
        "val_iou",
    ]
    val_cols = [c for c in display_cols if c.startswith("val_")]
    col_max: dict[str, float] = {}
    for col in val_cols:
        vals = []
        for row in rows:
            v = row.get(col, "")
            try:
                vals.append(float(v))
            except (TypeError, ValueError):
                continue
        col_max[col] = max(vals) if vals else float("nan")

    t = doc.add_table(rows=1, cols=len(display_cols))
    try:
        t.style = "Table Normal"
    except KeyError:
        try:
            t.style = "Normal Table"
        except KeyError:
            pass
    hdr = t.rows[0].cells
    for j, col in enumerate(display_cols):
        hdr[j].text = col.replace("val_", "val ").replace("_", " ").title()
        set_cell_shading(hdr[j], "E8E8E8")
        for p in hdr[j].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
                r.font.color.rgb = _black()
    for row in rows:
        cells = t.add_row().cells
        mid = (row.get("model_id") or "").lower()
        is_prop = mid == PROPOSED_MODEL_ID.lower()
        for j, col in enumerate(display_cols):
            v = row.get(col, "")
            if col.startswith("val_") and v not in ("", "nan", None):
                try:
                    fv = float(v)
                    if col in val_cols:
                        cells[j].text = f"{fv:.4f}"
                    else:
                        cells[j].text = str(v)
                    bold_cell = False
                    if is_prop and col in col_max and math.isfinite(col_max[col]):
                        if abs(fv - col_max[col]) <= 1e-6:
                            bold_cell = True
                    for p in cells[j].paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(8)
                            r.font.color.rgb = _black()
                            if bold_cell:
                                r.bold = True
                except ValueError:
                    cells[j].text = str(v)
            else:
                cells[j].text = str(v)
            for p in cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
                    r.font.color.rgb = _black()
    apply_journal_horizontal_rules_table(t)

    note = doc.add_paragraph()
    note.paragraph_format.left_indent = Inches(0.08)
    r0 = note.add_run("Note: ")
    r0.bold = True
    r0.font.size = Pt(9)
    r0.font.color.rgb = _black()
    r1 = note.add_run(
        "Values are best-epoch validation metrics from the unified ablation harness. "
        "Bold entries in the proposed-model row indicate column-wise maxima among all listed runs."
        + (f" {note_extra}" if note_extra else "")
    )
    r1.font.size = Pt(9)
    r1.font.color.rgb = _black()


NEW_ABSTRACT = (
    "Abstract. Landslides always pose a critical threat to both human life and infrastructure. "
    "With the advent of deep learning, CNNs, attention-based networks, and vision transformers have made "
    "significant progress in automating landslide segmentation from remote sensing data, yet problems persist, "
    "particularly in retaining local spatial features and adapting to new landscapes when models are trained "
    "only on a single modality or a rigid fusion rule. In this study we present TriEncoderCFMNet, an end-to-end "
    "segmentation architecture that reads RGB, topography (DEM with finite-difference slope cues), and a "
    "context stack aligned with competition-style multispectral conditioning. A trimodal gated fusion module "
    "combines pyramid features at each scale before a conditional flow-matching velocity field refines the "
    "mask latent, alongside an auxiliary convolutional head and geomorphology-aware regularisation. The model "
    "is evaluated on the Landslide4Sense benchmark with an extended baseline grid under a unified training "
    "protocol; Bijie metrics appear in Table 2 with the same ordering conventions. Compared with classical U-Net "
    "families and a strong dual-stream gated baseline, the proposed pipeline offers a different inductive bias—"
    "explicit three-way fusion plus flow-based refinement—while remaining comparable in training tooling to "
    "the rest of the ablation repository."
)

NEW_CONTRIBUTION_PARA = (
    "Motivated by these issues, the present contribution moves away from the earlier two-stage classifier–SAM "
    "pipeline and instead studies a single segmentation backbone where RGB, DEM, and context tensors are encoded "
    "in parallel, fused with data-dependent gates, and refined with a lightweight conditional flow matching decoder. "
    "We still care about the same practical worries—domain shift, ambiguous texture, and the need for topography—"
    "but we address them with trimodal fusion and flow-based denoising rather than promptable ViT adapters. "
    "The experimental section therefore reports Landslide4Sense validation metrics for a broad set of public "
    "baselines and for the new model under the same split and logging conventions as the rest of SAM/ablation_study."
)

METHOD_BODY = (
    "3.2 Proposed segmentation backbone (TriEncoderCFMNet). "
    "Let each training sample provide RGB image X_rgb, a DEM channel X_dem, and a context tensor X_ctx built from "
    "the competition stack (for Landslide4Sense this follows the six-channel context convention used elsewhere in "
    "the repository; Bijie uses a reduced four-channel stack). Each stream passes through the same pyramid encoder "
    "layout: a 7×4 stem, residual stages, and two stride-2 downsamples so that three scale maps A^(ℓ), B^(ℓ), C^(ℓ) "
    "are produced for ℓ ∈ {0,1,2}. Before the DEM encoder, finite-difference gradients build a slope magnitude map "
    "that is concatenated to elevation so the topography branch sees relief explicitly.\n\n"
    "At every scale the three feature maps are first mixed with a global softmax gate computed from their "
    "concatenation after global average pooling, yielding stream weights (w0,w1,w2) that sum to one. A spatial map "
    "γ^(ℓ) in [0,1] is predicted with a depthwise-heavy 7×7 path and used to blend the gated mixture with an "
    "unweighted average, followed by batch normalisation. The finest fused map feeds an auxiliary two-layer "
    "convolutional head whose logits are upsampled to full resolution.\n\n"
    "Conditional flow matching targets a smooth latent z obtained by logit-transforming the clipped mask. "
    "A velocity U-Net v_θ(x_t,t | fused pyramid) is trained with mean-squared error to the straight-path target "
    "velocity ε−z, together with a small finite-difference penalty that discourages rapid oscillations of v_θ in "
    "time. At validation one may integrate the learned field with a few Euler steps starting from Gaussian noise "
    "and add the resulting logits to the auxiliary head with a configurable scale. The segmentation objective "
    "combines Tversky loss on the auxiliary logits with a geomorphological term that down-weights mask-gradient "
    "energy on flat terrain so steep scarps can keep sharp boundaries. Optimisation follows the repository trainer: "
    "AdamW, ReduceLROnPlateau on a chosen validation score, gradient clipping, and epoch-wise CSV logging so "
    "figures in this manuscript can be reproduced from epoch_metrics.csv without ad-hoc post-processing.\n\n"
    "Relation to dual-stream baselines. The dual_stream_gated reference in this codebase fuses two EfficientNet "
    "towers before a UNet decoder. TriEncoderCFMNet differs by (i) using three parallel CNN pyramids instead of two "
    "ImageNet backbones, (ii) applying fusion at every pyramid level with both global softmax weights and a "
    "spatial γ map, and (iii) adding the flow-matching refinement path instead of relying on the classical decoder "
    "alone. Those design choices are orthogonal to the literature on gated fusion: they borrow the same intuition—"
    "that optical and elevation cues should not be averaged blindly—but implement it in a lighter stack suited to "
    "the ablation harness used here."
)

RESULTS_NARRATIVE = (
    "4 Results (Landslide4Sense). "
    "Table 1 summarises each architecture’s best validation epoch under the shared logging format exported by "
    "build_l4s_ablation_report.py. Metrics are pixel-level validation scores at the stored threshold defaults; "
    "AUROC and AUPRC come from the scalar columns logged each epoch, so the ROC/PR style figures in Fig. 7–12 "
    "trace these scalars against time with markers at epochs 24, 34, and 39 to mirror the layout of the earlier "
    "conference submission even though the underlying logs do not yet store full threshold sweeps.\n\n"
    "Fig. 2 and Fig. 5 show training and validation dynamics for the focus run (highest validation F1 in the table "
    "by default); Fig. 5 is a wide panel that couples the validation heatmap with training–versus–validation bars "
    "for precision, recall, and accuracy at the best epoch. Fig. 3 and Fig. 6 give the train–validation curve panels, "
    "while Fig. 4 overlays all models at their best epoch for F1, IoU, precision, and recall. "
    "Fig. 13 switches to a grouped precision–recall bar view across every model so the reader can compare the two "
    "operating points side by side without reading scatter coordinates.\n\n"
    "The supplementary overlay figures fig01–fig04 and fig06–fig09 in the same results folder provide additional cross-model "
    "comparisons (loss trajectories, ranked F1, radar-style summaries) that are not part of the main figure "
    "numbering but are useful for internal review."
)

DISCUSSION = (
    "5 Discussion. "
    "The table makes clear that classical baselines such as DeepLabV3+ and the dual_stream_gated configuration "
    "remain very strong on Landslide4Sense when training budget and augmentations are aligned. TriEncoderCFMNet "
    "should therefore be read less as a universal winner and more as a controlled experiment in richer fusion plus "
    "flow refinement: when the auxiliary head and FM terms are well balanced, the model trades a small amount of "
    "pixel F1 for higher AUROC in some checkpoints, which is consistent with smoother decision boundaries learned "
    "by the velocity field. Several baselines exhibit higher recall at the cost of precision; the grouped Fig. 13 "
    "makes that trade-off explicit for the whole cohort.\n\n"
    "Limitations include the absence of exported ROC/PR curve points per epoch, the mismatch in backbone capacity "
    "between EfficientNet-based dual-stream gated runs and the lighter CNN pyramids used here, and residual "
    "domain shift between the competition-style Landslide4Sense stack and the four-channel Bijie trimodal setup."
)

DEEP_ABLATION = (
    "5.1 Extended ablation discussion. "
    "LinkNet and the depth-aware U-Net variant sit at the lower end of validation F1 in this sweep, which matches "
    "their simpler inductive bias on heterogeneous Himalayan texture. GMNet and RMAU-Net climb toward the mid "
    "pack, suggesting that moderate architectural depth helps but does not by itself solve boundary ambiguity. "
    "TransUNet and EMR-HRNet introduce transformer or high-resolution pathways; their metrics benefit but still "
    "lag behind the strongest CNN decoders on this split, echoing findings elsewhere that ViT backbones can under-use "
    "very small landslide polygons without careful tuning. Shapeformer occupies a similar band, highlighting that "
    "structure-aware designs need dataset-specific schedules to shine.\n\n"
    "Dual-stream U-Net improves markedly once elevation is split into a parallel encoder, which is exactly the "
    "phenomenon emphasised in gated dual-stream papers: letting the network decide how much spectral versus "
    "topographic evidence to trust per scale is more stable than early fusion. The dual_stream_gated checkpoint "
    "included here uses the repository’s gated EfficientNet fusion before a UNet decoder; it remains one of the "
    "strongest baselines and is the fairest reference point when arguing about whether the extra CFM pathway in "
    "TriEncoderCFMNet pays for itself given similar training time.\n\n"
    "TriEncoder_cfm_v2 closes the narrative by showing where the new fusion and flow stack lands relative to that "
    "frontier. Interpreting the gap purely in terms of parameters would ignore optimisation noise, threshold choice, "
    "and the FM residual scaling; future work should log micro-F1 sweeps and full ROC tensors to tighten comparisons."
)

NEW_CONCLUSION = (
    "6 Conclusions. "
    "This manuscript retargets the earlier remote-sensing landslide study toward TriEncoderCFMNet, a trimodal "
    "gated fusion architecture with conditional flow matching and topography-aware losses. The introduction and "
    "related-work survey largely carry over because the underlying hazard remains the same even though the technical "
    "solution changed. Empirically, Landslide4Sense results in Table 1 and Figs. 2–13 together with Bijie results "
    "in Table 2 and Figs. 14–25 show how the new model sits inside a broad baseline lattice that includes dual-stream "
    "gated references inspired by the literature review PDF bundled with the repository. The final vector "
    "architecture diagram (Fig. 1) remains a placeholder so it can be swapped without renumbering downstream panels."
)

PLACEHOLDER_BIJIE = (
    "Bijie experiments use the repository `build_bijie_split` protocol (RGB composites with paired DEM and "
    "four-channel trimodal tensors for TriEncoderCFMNet) alongside the same baseline harness as Landslide4Sense. "
    "Table 2 and Figs. 14–25 summarise the dedicated Bijie summary CSV and the `conference_bijie` figure bundle."
)

BIJIE_RESULTS_CAPTION = (
    "4.2 Bijie benchmark. Table 2 lists the same model cohort on the Bijie split (best validation epoch per row). "
    "Figs. 14–25 repeat the conference-style diagnostics on the focus run `tri_encoder_cfm_v2` and the full Bijie "
    "leaderboard so that behaviour can be compared directly with the Landslide4Sense block above."
)

PLACEHOLDER_ARCH = (
    "[Placeholder — Fig. 1 architecture] "
    "Insert the final TriEncoderCFMNet diagram (export from draw.io, PowerPoint, or MODEL_ARCHITECTURE assets) "
    "here before submission."
)


def _nil_all_cell_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for edge in ("top", "bottom", "left", "right"):
                _set_cell_edge_border(cell, edge, "nil")


def add_numbered_equation_row(doc: Document, body: str, number: str) -> None:
    tbl = doc.add_table(rows=1, cols=2)
    try:
        tbl.style = "Table Normal"
    except KeyError:
        try:
            tbl.style = "Normal Table"
        except KeyError:
            pass
    _nil_all_cell_borders(tbl)
    left = tbl.rows[0].cells[0]
    right = tbl.rows[0].cells[1]
    pl = left.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl = pl.add_run(body)
    rl.font.size = Pt(10)
    rl.italic = True
    rl.font.color.rgb = _black()
    pr = right.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = pr.add_run(number)
    rr.font.size = Pt(10)
    rr.font.color.rgb = _black()
    doc.add_paragraph("")


def add_equations_section(doc: Document) -> None:
    doc.add_heading("3.2.1 Key equations", level=3)
    intro = doc.add_paragraph(
        "The following notation matches the trimodal backbone in §3.2; tensors omit explicit spatial indices."
    )
    for r in intro.runs:
        r.font.color.rgb = _black()
    equations = [
        (
            "Streams. For each pyramid level ℓ ∈ {0,1,2}, encoders emit aligned maps A(ℓ), B(ℓ), and C(ℓ) "
            "for RGB, topography (DEM with finite-difference slope cues), and the multispectral context stack.",
            "(1)",
        ),
        (
            "Global weights. Stacked global averages u(ℓ) pass through a linear map followed by softmax, "
            "yielding w(ℓ) ∈ ℝ³ with non-negative entries summing to one.",
            "(2)",
        ),
        (
            "Spatial gate and fusion. A depthwise-heavy head predicts γ(ℓ) ∈ [0,1] on the feature grid; "
            "fused tensors follow F(ℓ) = BN( γ(ℓ) ⊙ Σ_k w_k(ℓ) V_k(ℓ) + (1 − γ(ℓ)) ⊙ mean_k V_k(ℓ) ) "
            "after channel alignment of per-stream maps V_k.",
            "(3)",
        ),
        (
            "Auxiliary logits. Decoder layers on the finest fused map produce ŷ_aux at full mask resolution.",
            "(4)",
        ),
        (
            "Flow matching. With clipped mask M, latent z = logit(clip(M)), noise ε, interpolant x_t = (1−t)ε + t z, "
            "and velocity field v_θ, minimise 𝔼_{t,ε} ‖ v_θ(x_t, t | {F(ℓ)}) − (z − ε) ‖².",
            "(5)",
        ),
        (
            "Objective. Total training loss L = L_Tversky(ŷ_aux, M) + λ_FM L_FM + λ_geo L_geo with nonnegative weights.",
            "(6)",
        ),
    ]
    for body, num in equations:
        add_numbered_equation_row(doc, body, num)


def _add_algo_line_cell(cell, line: str, indent_in: float = 0.0) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent_in)
    parts = line.split("**")
    for i, chunk in enumerate(parts):
        if not chunk:
            continue
        r = p.add_run(chunk)
        r.bold = i % 2 == 1
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        r.font.color.rgb = _black()


def add_algorithm_training_block(doc: Document) -> None:
    doc.add_heading("3.2.2 Algorithm 1 — TriEncoderCFMNet training", level=3)
    outer = doc.add_table(rows=1, cols=1)
    try:
        outer.style = "Table Normal"
    except KeyError:
        outer.style = "Normal Table"
    cell = outer.rows[0].cells[0]
    _set_cell_edge_border(cell, "top", "single", "24")
    _set_cell_edge_border(cell, "bottom", "single", "24")
    _set_cell_edge_border(cell, "left", "nil")
    _set_cell_edge_border(cell, "right", "nil")

    pt = cell.paragraphs[0]
    t1 = pt.add_run("Algorithm 1")
    t1.bold = True
    t1.font.name = "Times New Roman"
    t1.font.size = Pt(11)
    t1.font.color.rgb = _black()
    t2 = pt.add_run("  TriEncoderCFMNet (trimodal fusion + CFM)")
    t2.font.name = "Times New Roman"
    t2.font.size = Pt(11)
    t2.font.color.rgb = _black()

    ps = cell.add_paragraph()
    s = ps.add_run(" " * 8)
    s.font.size = Pt(4)
    s.font.color.rgb = _black()

    algo = [
        ("**Inputs:** Dataset D of tuples (X_rgb, X_dem, X_ctx, M) with paired masks.", 0.0),
        ("**Parameters:** Weights θ over encoders, fusion heads, FM U-Net v_θ, auxiliary decoder; learning rate η; loss weights λ_FM, λ_geo.", 0.0),
        ("", 0.0),
        ("**for** each training epoch **do**", 0.0),
        ("    sample minibatch (X_rgb, X_dem, X_ctx, M) from D", 0.12),
        ("    forward encoders → {(A(ℓ), B(ℓ), C(ℓ))}; fuse → {F(ℓ)}; compute ŷ_aux", 0.12),
        ("    sample t ~ Uniform(0,1), ε ~ Normal(0, I); form x_t = (1−t)ε + t·logit(clip(M))", 0.12),
        ("    accumulate L_FM = ‖ v_θ(x_t, t | {F(ℓ)}) − (logit(clip(M)) − ε) ‖²", 0.12),
        ("    L = L_Tversky(ŷ_aux, M) + λ_FM · L_FM + λ_geo · L_geo", 0.12),
        ("    θ ← θ − η · ∇_θ L   (AdamW + gradient clipping as in trainer config)", 0.12),
        ("**end for**", 0.0),
        ("**return** θ", 0.0),
    ]
    for text, ind in algo:
        if text:
            _add_algo_line_cell(cell, text, ind)
    doc.add_paragraph("")


ORDERED_CONFERENCE_FIGURES: list[tuple[str, str]] = [
    ("Fig02_performance_heatmap_segmentation_model_training.png", "Performance heatmap for segmentation model (training dynamics)."),
    ("Fig03_train_validation_comparison_segmentation_model.png", "Train validation comparison for segmentation model."),
    ("Fig04_final_performance_summary_segmentation_model.png", "Final performance summary for segmentation model."),
    ("Fig05_performance_heatmap_segmentation_model_validation.png", "Validation heatmap with train–validation summary bars at best epoch (wide layout)."),
    ("Fig06_training_validation_comparison_segmentation_model.png", "Training validation comparison for segmentation model."),
    ("Fig07_ROC_curve_epoch_24.png", "ROC-style panel for epoch 24 (scalar AUROC vs epoch; full ROC pending sweep logs)."),
    ("Fig08_ROC_curve_epoch_34.png", "ROC-style panel for epoch 34."),
    ("Fig09_ROC_curve_epoch_39.png", "ROC-style panel for epoch 39."),
    ("Fig10_PR_curve_epoch_24.png", "PR-style panel for epoch 24 (scalar AUPRC vs epoch)."),
    ("Fig11_PR_curve_epoch_34.png", "PR-style panel for epoch 34."),
    ("Fig12_PR_curve_epoch_39.png", "PR-style panel for epoch 39."),
    ("Fig13_precision_recall_comparison_all_segmentation_models.png", "Precision–Recall comparison of all segmentation models (grouped bars)."),
]


def pack_submission_bundle(
    bundle_dir: Path,
    docx_src: Path,
    arch_md_src: Path,
    companion_md_src: Path,
    l4s_conf: Path,
    bijie_conf: Path,
    l4s_overlay_dir: Path,
    repo_root: Path,
    l4s_csv: Path,
    bijie_csv: Path,
    presence_report: Path,
) -> None:
    """Single flat folder: docx, MDs, summary CSVs, presence CSVs/PNG, prefixed conference PNGs."""
    if bundle_dir.exists():
        shutil.rmtree(bundle_dir)
    bundle_dir.mkdir(parents=True)

    shutil.copy2(docx_src, bundle_dir / docx_src.name)
    shutil.copy2(arch_md_src, bundle_dir / "TRI_ENCODER_MODEL_SPEC.md")
    if companion_md_src.is_file():
        shutil.copy2(companion_md_src, bundle_dir / companion_md_src.name)

    presence_md = repo_root / "SAM/resources/results/LANDSLIDE_PRESENCE_DETECTION.md"
    if not presence_md.is_file():
        presence_md = repo_root / "resources/results/LANDSLIDE_PRESENCE_DETECTION.md"
    if presence_md.is_file():
        shutil.copy2(presence_md, bundle_dir / "LANDSLIDE_PRESENCE_DETECTION.md")

    for csv_src in (l4s_csv, bijie_csv):
        if csv_src.is_file():
            shutil.copy2(csv_src, bundle_dir / csv_src.name)

    if presence_report.is_dir():
        for name in (
            "tri_encoder_presence_combined_table.csv",
            "tri_encoder_presence_run_manifest.csv",
            "tri_encoder_presence_images_bijie.csv",
            "tri_encoder_presence_images_l4s.csv",
        ):
            p = presence_report / name
            if p.is_file():
                shutil.copy2(p, bundle_dir / name)
        fig = presence_report / "fig_tri_encoder_presence_score_histogram.png"
        if fig.is_file():
            shutil.copy2(fig, bundle_dir / fig.name)

    if l4s_conf.is_dir():
        for p in sorted(l4s_conf.glob("Fig*.png")) + sorted(l4s_conf.glob("Fig*.txt")):
            shutil.copy2(p, bundle_dir / f"l4s_{p.name}")
    if bijie_conf.is_dir():
        for p in sorted(bijie_conf.glob("Fig*.png")) + sorted(bijie_conf.glob("Fig*.txt")):
            shutil.copy2(p, bundle_dir / f"bijie_{p.name}")
    if l4s_overlay_dir.is_dir():
        for i in range(1, 10):
            if i == 5:
                continue
            for p in sorted(l4s_overlay_dir.glob(f"fig{i:02d}_*.png")):
                shutil.copy2(p, bundle_dir / f"overlay_{p.name}")


def main() -> None:
    root = repo_root()
    pdf_conf = root / "SAM/resources/docs/conference_remotesensing_landslide.pdf"
    md_path = root / "SAM/model_architecture_cfm_landseg/MODEL_ARCHITECTURE.md"
    csv_path = root / "SAM/resources/results/l4s_ablation_report/landslide4sense_best_validation_summary.csv"
    bijie_csv = root / "SAM/resources/results/bijie_ablation_report/bijie_best_validation_summary.csv"
    fig_dir = root / "SAM/resources/results/l4s_ablation_report/paper_comparison_figures"
    conf_fig = fig_dir / "conference_remotesensing_landslide"
    bijie_conf_fig = (
        root / "SAM/resources/results/bijie_ablation_report/paper_comparison_figures/conference_bijie"
    )
    companion_md = root / "SAM/resources/results/bijie_ablation_report/BIJIE_AND_L4S_ABLATION_COMPANION.md"
    bundle_dir = root / "SAM/model_architecture_cfm_landseg/paper_submission_bundle"
    out_docx = root / "SAM/model_architecture_cfm_landseg/conference_manuscript_tri_encoder_cfm.docx"

    if not pdf_conf.is_file():
        raise SystemExit(f"Missing conference PDF: {pdf_conf}")

    full = clean_header_lines(extract_pdf_text(pdf_conf))

    title_block = slice_between(full, r"^Cross-Region Landslide", r"Samreedh", re.MULTILINE)
    # title might be broken - manual title
    authors = "Samreedh Bhuyan, Debasish Saikia, and Hridoy Jyoti Mahanta"

    intro_raw = slice_between(full, r"1 Introduction", r"2 Related Works")
    intro = strip_first_numbered_heading(intro_raw)
    intro = minimal_intro_edits(intro)

    related_raw = slice_between(full, r"2 Related Works", r"3 Materials and methods")
    related = strip_first_numbered_heading(related_raw)

    datasets_blurb = slice_between(
        full,
        r"3 Materials and methods\s+To ensure",
        r"3\.1 Dataset preprocessing",
    )
    datasets_blurb = strip_first_numbered_heading(datasets_blurb)
    preprocess = slice_between(
        full,
        r"3\.1 Dataset preprocessing",
        r"3\.2 Detailed workflow",
    )
    preprocess = strip_first_numbered_heading(preprocess)

    references = slice_between(full, r"References\s*\n", None)
    # Trim trailing page noise
    references = re.split(r"\n\d+\s+S\.\s+Bhuyan", references)[0].strip()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    h0 = doc.add_heading(
        "Cross-Region Landslide Segmentation with TriEncoderCFMNet "
        "(Trimodal Gated Fusion and Conditional Flow Matching)",
        level=0,
    )
    h0.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_auth = doc.add_paragraph(authors)
    p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    doc.add_heading("Abstract", level=1)
    add_para(doc, NEW_ABSTRACT)
    doc.add_paragraph("Keywords: Landslide modeling · Machine learning · Segmentation · "
                      "Topography fusion · Conditional flow matching · Remote sensing.")

    doc.add_heading("1 Introduction", level=1)
    add_para(doc, intro)

    doc.add_heading("2 Related works", level=1)
    add_para(doc, related)

    doc.add_heading("3 Materials and methods", level=1)
    doc.add_heading("3.1 Datasets and preprocessing", level=2)
    if datasets_blurb:
        add_para(doc, datasets_blurb)
    if preprocess:
        add_para(doc, preprocess)

    doc.add_heading("3.2 Proposed model", level=2)
    add_para(doc, METHOD_BODY)
    add_equations_section(doc)
    add_algorithm_training_block(doc)

    doc.add_paragraph("")
    p_ph = doc.add_paragraph(PLACEHOLDER_ARCH)
    for r in p_ph.runs:
        r.font.color.rgb = _black()
        r.bold = True

    doc.add_heading("3.3 Bijie dataset and protocol", level=2)
    p_b = doc.add_paragraph(PLACEHOLDER_BIJIE)
    for r in p_b.runs:
        r.font.color.rgb = _black()
        r.bold = True

    doc.add_heading("4 Results and figures", level=1)
    cap_tab = (
        "Table 1. Landslide4Sense best-epoch validation summary across segmentation models "
        "(unified ablation harness). Values come from landslide4sense_best_validation_summary.csv."
    )
    p_cap = doc.add_paragraph(cap_tab)
    p_cap.runs[0].bold = True
    add_results_table(
        doc,
        csv_path,
        note_extra="Dataset: Landslide4Sense benchmark.",
    )
    doc.add_paragraph("")

    add_para(doc, RESULTS_NARRATIVE)

    # Fig 1 placeholder box
    doc.add_heading("Figure 1 (placeholder)", level=2)
    add_para(doc, PLACEHOLDER_ARCH)

    for idx, (fname, rest) in enumerate(ORDERED_CONFERENCE_FIGURES, start=2):
        cap = f"Fig. {idx}. {rest}"
        w = 6.3
        if "Fig02" in fname or "Fig05" in fname:
            w = 7.05
        elif "Fig13" in fname:
            w = 6.55
        add_figure(doc, conf_fig / fname, cap, width_in=w)

    if bijie_csv.is_file() and bijie_conf_fig.is_dir():
        doc.add_heading("4.2 Bijie benchmark results", level=2)
        add_para(doc, BIJIE_RESULTS_CAPTION)
        p_cap2 = doc.add_paragraph(
            "Table 2. Bijie best-epoch validation summary across the same model cohort "
            "(values from bijie_best_validation_summary.csv)."
        )
        p_cap2.runs[0].bold = True
        add_results_table(
            doc,
            bijie_csv,
            note_extra="Dataset: Bijie City landslide split; AUROC/AUPRC absent where logs omit those scalars.",
        )
        doc.add_paragraph("")
        for i, (fname, rest) in enumerate(ORDERED_CONFERENCE_FIGURES, start=14):
            cap = f"Fig. {i} (Bijie). {rest}"
            w = 6.3
            if "Fig02" in fname or "Fig05" in fname:
                w = 7.05
            elif "Fig13" in fname:
                w = 6.55
            add_figure(doc, bijie_conf_fig / fname, cap, width_in=w)

    doc.add_heading("Supplementary overlay figures (fig01–fig04, fig06–fig09)", level=2)
    for i in range(1, 10):
        if i == 5:
            continue
        matches = sorted(fig_dir.glob(f"fig{i:02d}_*.png"))
        if not matches:
            continue
        add_figure(doc, matches[0], f"Overlay figure fig{i:02d}: {matches[0].stem}.", width_in=6.4)

    doc.add_heading("5 Discussion and extended ablation", level=1)
    add_para(doc, DISCUSSION)
    add_para(doc, DEEP_ABLATION)

    doc.add_heading("6 Conclusions", level=1)
    add_para(doc, NEW_CONCLUSION)

    doc.add_heading("References", level=1)
    add_para(doc, references, style=None)

    doc.add_page_break()
    doc.add_heading("Appendix A — Technical specification pointer", level=1)
    add_para(
        doc,
        f"Full equations, CLI defaults, and file-level pointers are documented in {md_path.relative_to(root)}.",
    )

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    force_all_text_black(doc)
    doc.save(str(out_docx))
    print(f"Wrote {out_docx}")
    presence_report = root / "SAM/resources/results/landslide_presence_report"
    if not presence_report.is_dir():
        presence_report = root / "resources/results/landslide_presence_report"
    pack_submission_bundle(
        bundle_dir,
        out_docx,
        md_path,
        companion_md,
        conf_fig,
        bijie_conf_fig,
        fig_dir,
        root,
        csv_path,
        bijie_csv,
        presence_report,
    )
    print(f"Packed submission bundle: {bundle_dir}")


if __name__ == "__main__":
    main()
